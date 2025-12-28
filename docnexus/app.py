"""
Markdown Documentation Viewer
A Flask-based web application that presents Markdown files from a folder as well-formatted HTML sections.
"""

from flask import Flask, render_template, send_from_directory, abort, request, make_response, session, jsonify
import os
import sys
import markdown
from pathlib import Path
from datetime import datetime
import io
import re
import html as html_module
import json
import shutil
import logging
from logging.handlers import RotatingFileHandler
import zipfile
import urllib.request
import subprocess
import tempfile
from bs4 import BeautifulSoup

# Try to import PDF libraries
PDF_EXPORT_AVAILABLE = False
PDF_LIB_NAME = None

try:
    import pdfkit
    PDF_EXPORT_AVAILABLE = True
    PDF_LIB_NAME = "pdfkit"
except Exception as e:
    print(f"Warning: pdfkit not available: {e}")
    print("PDF export feature will be disabled.")
    pdfkit = None

# Try to import Word export libraries
WORD_EXPORT_AVAILABLE = False
WORD_LIB_NAME = None

try:
    from htmldocx import HtmlToDocx
    WORD_EXPORT_AVAILABLE = True
    WORD_LIB_NAME = "htmldocx"
except Exception as e:
    print(f"Warning: htmldocx not available: {e}")
    print("Word export feature will be disabled.")
    HtmlToDocx = None

# Try to import Word input libraries
WORD_INPUT_AVAILABLE = False
try:
    import mammoth
    WORD_INPUT_AVAILABLE = True
except Exception as e:
    print(f"Warning: mammoth not available: {e}")
    print("Word input feature will be disabled.")
    mammoth = None

try:
    from docnexus.core.renderer import render_baseline, run_pipeline
    from docnexus.features.registry import FeatureManager, Feature, FeatureState
    from docnexus.features import smart_convert as smart
    from docnexus.features.standard import normalize_headings, sanitize_attr_tokens, build_toc, annotate_blocks
except Exception:
    # allow running as a script: add project root to sys.path, then absolute imports
    PROJECT_ROOT_FOR_PATH = Path(__file__).resolve().parent.parent
    if str(PROJECT_ROOT_FOR_PATH) not in sys.path:
        sys.path.insert(0, str(PROJECT_ROOT_FOR_PATH))
    from docnexus.core.renderer import render_baseline, run_pipeline
    from docnexus.features.registry import FeatureManager, Feature, FeatureState
    from docnexus.features import smart_convert as smart
    from docnexus.features.standard import normalize_headings, sanitize_attr_tokens, build_toc, annotate_blocks

try:
    from . import __version__
    VERSION = __version__
except ImportError:
    # Fallback if running directly
    VERSION = '1.0.1'

app = Flask(__name__, static_folder='static')
app.config['TEMPLATES_AUTO_RELOAD'] = True
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
app.secret_key = os.urandom(24)  # For session management
# Note: We do NOT set MAX_CONTENT_LENGTH here because:
# 1. Form-encoded data can be 2-3x larger than actual file content
# 2. We validate actual file/content size at the application level instead
# 3. This allows the server to accept the HTTP request and check the actual data size intelligently

# Configuration
# Get the project root directory (parent of doc_viewer)
# Configuration
# Get the project root directory (parent of doc_viewer)
if getattr(sys, 'frozen', False):
    # Running as compiled executable
    PROJECT_ROOT = Path(sys.executable).parent
else:
    # Running from source
    PROJECT_ROOT = Path(__file__).parent.parent

# Logging Configuration
LOG_DIR = PROJECT_ROOT / 'logs'
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / 'docnexus.log'

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        RotatingFileHandler(LOG_FILE, maxBytes=10*1024*1024, backupCount=7),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('docnexus')
logger.info(f"Application starting - Version {VERSION}")

# Workspace Configuration
CONFIG_FILE = PROJECT_ROOT / 'config.json'

def load_config():
    """Load workspace configuration."""
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Failed to load config: {e}")
    # Determine default workspace
    default_workspace = PROJECT_ROOT / 'workspace'
    if not default_workspace.exists() and (PROJECT_ROOT / 'examples').exists():
        default_workspace = PROJECT_ROOT / 'examples'

    return {
        'workspaces': [str(default_workspace)],
        'active_workspace': str(default_workspace),
        'recent_workspaces': []
    }

def save_config(config):
    """Save workspace configuration."""
    try:
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f, indent=2)
        logger.info("Configuration saved successfully")
    except Exception as e:
        logger.error(f"Failed to save config: {e}")

# Load configuration
CONFIG = load_config()
MD_FOLDER = Path(CONFIG['active_workspace'])  # Folder containing documents
DOCS_FOLDER = PROJECT_ROOT / 'docs'  # Documentation folder
ALLOWED_EXTENSIONS = {'.md', '.markdown', '.txt', '.docx'}

# File size limits (in bytes)
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB for actual file content
MAX_EXPORT_HTML_SIZE = 50 * 1024 * 1024  # 50 MB for export HTML content

# Feature registry: keep smart/experimental separate from baseline rendering
FEATURES = FeatureManager()
# STANDARD features (always run)
FEATURES.register(Feature("STD_NORMALIZE", normalize_headings, FeatureState.STANDARD))
FEATURES.register(Feature("STD_SANITIZE_ATTR", sanitize_attr_tokens, FeatureState.STANDARD))
FEATURES.register(Feature("STD_TOC", build_toc, FeatureState.STANDARD))
FEATURES.register(Feature("STD_ANNOTATE", annotate_blocks, FeatureState.STANDARD))
# EXPERIMENTAL features (smart toggle) - Registered but disabled in v1.0.0
# Will be re-enabled in v1.1/1.2 with proper UI - see doc/FUTURE_FEATURES.md
FEATURES.register(Feature("SMART_TABLES", smart.convert_ascii_tables_to_markdown, FeatureState.EXPERIMENTAL))
FEATURES.register(Feature("SMART_SIP", smart.convert_sip_signaling_to_mermaid, FeatureState.EXPERIMENTAL))
FEATURES.register(Feature("SMART_TOPOLOGY", smart.convert_topology_to_mermaid, FeatureState.EXPERIMENTAL))

# Context Processor for Debugging (moved here after all config is loaded)
@app.context_processor
def inject_debug_info():
    try:
        # Resolve template folder to absolute path
        template_path = Path(app.template_folder)
        if not template_path.is_absolute():
            template_path = Path(app.root_path) / template_path
        
        return {
            'debug_info': {
                'project_root': str(PROJECT_ROOT),
                'md_folder': str(MD_FOLDER),
                'template_folder': str(template_path),
                'frozen': getattr(sys, 'frozen', False),
                'python_path': str(sys.executable),
                'version': VERSION
            }
        }
    except Exception as e:
        logger.error(f"Error in context processor: {e}")
        return {'debug_info': {'error': str(e)}}

# SIP Protocol Knowledge Base
SIP_KNOWLEDGE = {
    'request_methods': [
        'INVITE', 'ACK', 'BYE', 'CANCEL', 'REGISTER', 'OPTIONS',
        'PRACK', 'SUBSCRIBE', 'NOTIFY', 'PUBLISH', 'INFO', 'REFER',
        'MESSAGE', 'UPDATE'
    ],
    'response_codes': {
        # 1xx Provisional
        '100': 'Trying', '180': 'Ringing', '181': 'Call Is Being Forwarded',
        '182': 'Queued', '183': 'Session Progress',
        # 2xx Success
        '200': 'OK', '202': 'Accepted',
        # 3xx Redirection
        '300': 'Multiple Choices', '301': 'Moved Permanently', '302': 'Moved Temporarily',
        # 4xx Client Error
        '400': 'Bad Request', '401': 'Unauthorized', '403': 'Forbidden',
        '404': 'Not Found', '408': 'Request Timeout', '486': 'Busy Here',
        '487': 'Request Terminated',
        # 5xx Server Error
        '500': 'Server Internal Error', '503': 'Service Unavailable',
        # 6xx Global Failure
        '600': 'Busy Everywhere', '603': 'Decline', '604': 'Does Not Exist Anywhere'
    },
    'call_flow_patterns': {
        'basic_call': ['INVITE', '100 Trying', '180 Ringing', '200 OK', 'ACK', 'BYE', '200 OK'],
        'with_prack': ['INVITE', '100 Trying', '183 Session Progress', 'PRACK', '200 OK', '180 Ringing', 'PRACK', '200 OK', '200 OK', 'ACK'],
        'cancel': ['INVITE', '100 Trying', 'CANCEL', '200 OK', '487 Request Terminated', 'ACK'],
        'busy': ['INVITE', '100 Trying', '486 Busy Here', 'ACK'],
        'declined': ['INVITE', '100 Trying', '603 Decline', 'ACK']
    },
    'headers': [
        'Via', 'From', 'To', 'Call-ID', 'CSeq', 'Contact', 'Max-Forwards',
        'Content-Type', 'Content-Length', 'User-Agent', 'Allow', 'Supported'
    ],
    'sdp_attributes': ['RTP', 'SRTP', 'RTCP', 'codec', 'sendrecv', 'recvonly', 'sendonly']
}

def get_markdown_files():
    """Get all markdown files from the specified folder and subfolders."""
    md_path = Path(MD_FOLDER)
    if not md_path.exists():
        md_path.mkdir(parents=True, exist_ok=True)
        return []
    
    files = []
    # Use rglob to recursively find all markdown files
    for file_path in md_path.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in ALLOWED_EXTENSIONS:
            stat = file_path.stat()
            # Get relative path from doc_gallery folder
            rel_path = file_path.relative_to(md_path)
            folder = str(rel_path.parent) if rel_path.parent != Path('.') else ''
            
            # Format size
            size_bytes = stat.st_size
            if size_bytes < 1024:
                size_str = f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                size_str = f"{size_bytes / 1024:.1f} KB"
            else:
                size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
            
            # Keep raw timestamp for sorting
            raw_modified = stat.st_mtime

            files.append({
                'name': file_path.stem,
                'filename': file_path.name,
                'type': file_path.suffix.lstrip('.').lower(),
                'path': file_path,
                'relative_path': str(rel_path).replace('\\', '/'),
                'folder': folder.replace('\\', '/') if folder else '',
                'size': size_str,
                'raw_size': size_bytes,
                'modified': datetime.fromtimestamp(raw_modified).strftime('%Y-%m-%d %H:%M'),
                'raw_modified': raw_modified
            })
    
    # Sort by folder, then by modification time (newest first)
    files.sort(key=lambda x: (x['folder'], -x['raw_modified']))
    return files

def convert_ascii_tables_to_markdown(content):
    """Convert ASCII tables, network topology, and SIP flow diagrams to appropriate formats."""
    import re
    
    def get_preceding_heading(content, code_block_start):
        """Get the heading that precedes a code block."""
        lines_before = content[:code_block_start].split('\n')
        # Look backwards for the nearest heading
        for line in reversed(lines_before[-10:]):  # Check last 10 lines
            if line.startswith('#'):
                return line.strip()
        return ''
    
    def should_convert_topology(heading):
        """Check if heading suggests this should be a topology diagram."""
        topology_keywords = ['topology', 'network', 'architecture', 'setup', 'configuration']
        heading_lower = heading.lower()
        return any(keyword in heading_lower for keyword in topology_keywords)
    
    def should_convert_signaling(heading):
        """Check if heading suggests this should be a signaling/flow diagram."""
        flow_keywords = ['flow', 'signaling', 'sequence', 'call flow', 'message flow', 'sip']
        heading_lower = heading.lower()
        return any(keyword in heading_lower for keyword in flow_keywords)
    
    def detect_network_topology(lines):
        """Detect if this is a network topology diagram."""
        text = '\n'.join(lines)
        # Check for topology indicators
        topology_indicators = ['UAC', 'UAS', 'Router', 'Server', 'Switch', 'Gateway', 'Proxy']
        box_chars = ['┌', '─', '┐', '│', '└', '┘', '├', '┤', '┬', '┴']
        
        has_boxes = sum(1 for char in box_chars if char in text) >= 5
        has_topology_terms = sum(1 for term in topology_indicators if term in text) >= 2
        has_ip_addresses = bool(re.search(r'\d+\.\d+\.\d+\.\d+', text))
        
        # Must have boxes AND either topology terms or IP addresses
        return has_boxes and (has_topology_terms or has_ip_addresses)
    
    def detect_sip_signaling(lines):
        """Detect if this is a SIP signaling flow diagram using SIP knowledge."""
        text = '\n'.join(lines)
        text_upper = text.upper()
        
        # Check for SIP request methods
        request_count = sum(1 for method in SIP_KNOWLEDGE['request_methods'] if method in text_upper)
        
        # Check for SIP response codes
        response_count = 0
        for code, desc in SIP_KNOWLEDGE['response_codes'].items():
            if f'{code} {desc}' in text or code in text or desc.upper() in text_upper:
                response_count += 1
        
        # Flow indicators
        has_arrows = any(arrow in text for arrow in ['──>', '<──', '→', '←', '────>', '<────'])
        has_timing = 'T+' in text or 'TIME' in lines[0] if lines else False
        has_process_notes = any(char in text for char in ['┌', '├', '└'])
        
        # Strong indicators: arrows + (timing OR process notes) + (requests OR responses)
        return has_arrows and (has_timing or has_process_notes) and (request_count >= 2 or response_count >= 2)
    
    def convert_topology_to_mermaid(lines):
        """Convert network topology to Mermaid flowchart."""
        text = '\n'.join(lines)
        
        # Extract network elements
        patterns = [
            r'(UAC|UAS)\s*\(([^)]+)\)',
            r'(Crestone Router|Router|Server|Gateway|Proxy)\s*(\d+)?\s*\(([^)]+)\)',
            r'(Crestone Router \d+)',
        ]
        
        nodes = []
        node_details = {}
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                full_match = match.group(0)
                # Extract name and details
                if '(' in full_match:
                    name_part = full_match.split('(')[0].strip()
                    detail_part = full_match.split('(')[1].split(')')[0].strip()
                    if name_part not in nodes:
                        nodes.append(name_part)
                        node_details[name_part] = detail_part
                else:
                    if full_match not in nodes and len(full_match) < 50:
                        nodes.append(full_match)
        
        # If we found nodes, create a flowchart
        if len(nodes) >= 2:
            mermaid_lines = ['```mermaid', 'flowchart LR']
            
            # Create nodes with details
            for i, node in enumerate(nodes):
                node_id = f'N{i}'
                if node in node_details:
                    label = f"{node}<br/>{node_details[node]}"
                else:
                    label = node
                mermaid_lines.append(f'    {node_id}["{label}"]')
            
            # Create connections (left-to-right flow)
            for i in range(len(nodes) - 1):
                mermaid_lines.append(f'    N{i} -->|TCP/RTP| N{i+1}')
            
            # Style
            mermaid_lines.append('    classDef default fill:#e3f2fd,stroke:#1976d2,stroke-width:2px,color:#000')
            mermaid_lines.append('```')
            
            return '\n'.join(mermaid_lines)
        
        return None
    
    def convert_sip_signaling_to_mermaid(lines):
        """Convert SIP signaling flow to Mermaid sequence diagram using SIP knowledge."""
        # Extract participants from header
        participants = []
        
        # Find the line with participant names
        for i, line in enumerate(lines[:5]):
            if not line.strip().startswith('-') and 'Time' not in line:
                parts = re.split(r'\s{2,}', line.strip())
                if len(parts) >= 3:
                    participants = [p.strip() for p in parts if p.strip() and not p.startswith('T+')]
                    break
        
        # Clean participant names (remove IP addresses)
        clean_participants = []
        for p in participants:
            clean_p = re.sub(r'\s*\([^)]*\)', '', p).strip()
            if clean_p and clean_p.lower() not in ['time']:
                clean_participants.append(clean_p)
        
        participants = clean_participants[:5]
        
        if len(participants) < 2:
            participants = ['UAC', 'Router1', 'Router2', 'UAS']
        
        # Build Mermaid sequence diagram
        mermaid_lines = ['```mermaid', 'sequenceDiagram']
        
        # Add participants
        for p in participants:
            safe_name = re.sub(r'[^a-zA-Z0-9_]', '_', p)
            mermaid_lines.append(f'    participant {safe_name} as {p}')
        
        # Parse message flows using SIP knowledge
        current_time = None
        for line in lines[2:]:  # Skip header lines
            if line.strip().startswith('---'):
                continue
            
            # Extract timing
            time_match = re.match(r'(T\+\d+\w+)', line)
            if time_match:
                current_time = time_match.group(1)
            
            # Right arrow messages
            if '──>' in line or '───>' in line or '────>' in line or '→' in line:
                message = 'Message'
                
                # Check SIP request methods
                for method in SIP_KNOWLEDGE['request_methods']:
                    if method in line.upper():
                        message = method
                        if current_time:
                            message = f'{method} [{current_time}]'
                        break
                
                # Check SIP response codes
                if message == 'Message':
                    for code, desc in SIP_KNOWLEDGE['response_codes'].items():
                        if f'{code} {desc}' in line or f'{code}' in line[:10]:
                            message = f'{code} {desc}'
                            if current_time:
                                message = f'{code} {desc} [{current_time}]'
                            break
                
                # Extract additional info from line (RTP, SRTP, etc.)
                for attr in SIP_KNOWLEDGE['sdp_attributes']:
                    if attr in line and attr not in message:
                        message = f'{message} ({attr})'
                        break
                
                if len(participants) >= 2:
                    for i in range(len(participants) - 1):
                        src = re.sub(r'[^a-zA-Z0-9_]', '_', participants[i])
                        dst = re.sub(r'[^a-zA-Z0-9_]', '_', participants[i + 1])
                        mermaid_lines.append(f'    {src}->>{dst}: {message}')
                        break
            
            # Left arrow responses
            elif '<──' in line or '<───' in line or '<────' in line or '←' in line:
                message = 'Response'
                
                # Check SIP response codes
                for code, desc in SIP_KNOWLEDGE['response_codes'].items():
                    if f'{code} {desc}' in line or f'{code}' in line[:20]:
                        message = f'{code} {desc}'
                        if current_time:
                            message = f'{code} {desc} [{current_time}]'
                        break
                
                # Check for request methods in responses
                if message == 'Response':
                    for method in SIP_KNOWLEDGE['request_methods']:
                        if method in line.upper():
                            message = f'200 OK ({method})'
                            break
                
                if len(participants) >= 2:
                    for i in range(len(participants) - 1, 0, -1):
                        src = re.sub(r'[^a-zA-Z0-9_]', '_', participants[i])
                        dst = re.sub(r'[^a-zA-Z0-9_]', '_', participants[i - 1])
                        mermaid_lines.append(f'    {src}-->>{dst}: {message}')
                        break
            
            # Process notes
            elif any(char in line for char in ['┌', '├', '└']):
                note_text = re.sub(r'[┌├└─]', '', line).strip()
                if note_text and len(note_text) > 3 and len(participants) >= 2:
                    participant = re.sub(r'[^a-zA-Z0-9_]', '_', participants[1])
                    # Limit note length
                    if len(note_text) > 50:
                        note_text = note_text[:47] + '...'
                    mermaid_lines.append(f'    Note over {participant}: {note_text}')
        
        mermaid_lines.append('```')
        return '\n'.join(mermaid_lines)
    
    def detect_simple_table(lines):
        """Detect if this is a simple data table."""
        potential_table = []
        for line in lines:
            cols = re.split(r'\s{2,}', line.strip())
            if len(cols) > 1:
                potential_table.append(cols)
        
        if len(potential_table) >= 2:
            col_counts = [len(row) for row in potential_table]
            # Check for consistent columns
            return max(col_counts) - min(col_counts) <= 1
        return False
    
    def convert_table_to_markdown(lines):
        """Convert ASCII table to markdown table."""
        potential_table = []
        for line in lines:
            cols = re.split(r'\s{2,}', line.strip())
            if len(cols) > 1:
                potential_table.append(cols)
        
        max_cols = max(len(row) for row in potential_table)
        table_lines = []
        
        # Header
        header = potential_table[0]
        while len(header) < max_cols:
            header.append('')
        table_lines.append('| ' + ' | '.join(header) + ' |')
        
        # Separator
        table_lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
        
        # Data rows
        for row in potential_table[1:]:
            while len(row) < max_cols:
                row.append('')
            table_lines.append('| ' + ' | '.join(row) + ' |')
        
        return '\n' + '\n'.join(table_lines) + '\n'
    
    def process_code_block_with_context(match, preceding_heading):
        code_content = match.group(1)
        lines = code_content.strip().split('\n')
        
        if len(lines) < 2:
            return match.group(0)
        
        # Use heading context to guide conversion
        heading_suggests_topology = should_convert_topology(preceding_heading)
        heading_suggests_signaling = should_convert_signaling(preceding_heading)
        
        # Priority 1: Heading suggests topology + content matches
        if heading_suggests_topology and detect_network_topology(lines):
            converted = convert_topology_to_mermaid(lines)
            if converted:
                return '\n' + converted + '\n'
        
        # Priority 2: Heading suggests signaling + content matches  
        if heading_suggests_signaling and detect_sip_signaling(lines):
            return '\n' + convert_sip_signaling_to_mermaid(lines) + '\n'
        
        # Priority 3: Auto-detect without heading context (stricter criteria)
        if not heading_suggests_topology and not heading_suggests_signaling:
            # Only convert if very strong signals
            if detect_network_topology(lines):
                converted = convert_topology_to_mermaid(lines)
                if converted:
                    return '\n' + converted + '\n'
            
            if detect_sip_signaling(lines):
                return '\n' + convert_sip_signaling_to_mermaid(lines) + '\n'
        
        # Priority 4: Check for simple table
        if detect_simple_table(lines):
            return convert_table_to_markdown(lines)
        
        # Keep original
        return match.group(0)
    
    # Process code blocks with context awareness
    result = []
    last_end = 0
    
    for match in re.finditer(r'```\n(.*?)\n```', content, flags=re.DOTALL):
        # Add content before this code block
        result.append(content[last_end:match.start()])
        
        # Get preceding heading
        preceding_heading = get_preceding_heading(content, match.start())
        
        # Process the code block with context
        processed = process_code_block_with_context(match, preceding_heading)
        result.append(processed)
        
        last_end = match.end()
    
    # Add remaining content
    result.append(content[last_end:])
    
    return ''.join(result)

# ============================================================================
# HELPER FUNCTIONS FOR NEW FEATURES (v1.4.0)
# ============================================================================

def convert_docx_to_html(docx_path: Path) -> str:
    """Convert Word document to HTML using mammoth."""
    if not WORD_INPUT_AVAILABLE:
        raise Exception("mammoth library not available")
    
    logger.info(f"Converting Word document: {docx_path}")
    try:
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
            
            # Log any messages/warnings from conversion
            if result.messages:
                for msg in result.messages:
                    logger.warning(f"Mammoth conversion message: {msg}")
            
            logger.info(f"Successfully converted Word document, size: {len(html_content)} bytes")
            return html_content
    except Exception as e:
        logger.error(f"Failed to convert Word document: {e}", exc_info=True)
        raise

def process_links_in_html(html_content: str, base_path: Path = None, is_preview: bool = False) -> str:
    """
    Process all links in HTML to ensure they are clickable and properly resolved.
    - External links open in new tab
    - Relative links resolved based on document location
    """
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for a_tag in soup.find_all('a'):
            href = a_tag.get('href', '')
            
            if not href:
                continue
            
            # External links - open in new tab
            if href.startswith('http://') or href.startswith('https://'):
                a_tag['target'] = '_blank'
                a_tag['rel'] = 'noopener noreferrer'
                logger.debug(f"External link: {href}")
            
            # Email links
            elif href.startswith('mailto:'):
                pass  # Already functional
            
            # Anchor links within document
            elif href.startswith('#'):
                pass  # Work as-is
            
            # Absolute paths
            elif href.startswith('/'):
                pass  # Already absolute
            
            # Relative links - resolve based on document location
            else:
                if base_path:
                    try:
                        # Resolve relative to document's directory
                        resolved = (base_path / href).resolve()
                        
                        # Check if it exists and is within MD_FOLDER
                        if resolved.exists() and resolved.is_relative_to(MD_FOLDER):
                            rel_path = resolved.relative_to(MD_FOLDER)
                            a_tag['href'] = f'/file/{rel_path}'
                            logger.debug(f"Resolved relative link {href} -> /file/{rel_path}")
                        else:
                            # Link target doesn't exist or outside workspace
                            a_tag['class'] = (a_tag.get('class', []) or []) + ['broken-link']
                            a_tag['title'] = 'Link target not found'
                            a_tag['style'] = 'color: #dc2626; text-decoration: underline dotted;'
                            logger.warning(f"Broken link: {href} (resolved to {resolved})")
                    except Exception as e:
                        logger.warning(f"Failed to resolve link {href}: {e}")
        
        return str(soup)
    except Exception as e:
        logger.error(f"Error processing links: {e}", exc_info=True)
        return html_content  # Return original if processing fails

def is_safe_workspace(path: Path) -> bool:
    """Check if directory is safe to use as workspace."""
    try:
        path = path.resolve()
        
        # Blocked directories (Windows)
        blocked = [
            Path('C:\\Windows'),
            Path('C:\\Program Files'),
            Path('C:\\Program Files (x86)'),
            Path.home() / 'AppData',
        ]
        
        # Check against blocked paths
        for blocked_path in blocked:
            try:
                if path == blocked_path or path.is_relative_to(blocked_path):
                    return False
            except:
                pass
        
        # Must have read permission
        try:
            list(path.iterdir())  # Test read access
        except PermissionError:
            return False
        
        return True
    except Exception as e:
        logger.error(f"Error checking workspace safety: {e}")
        return False

def sanitize_log_content(content: str) -> str:
    """Remove sensitive information from logs."""
    # Remove IP addresses
    content = re.sub(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b', '<IP_ADDRESS>', content)
    # Remove full file paths - keep only filenames
    content = re.sub(r'[A-Z]:\\[^\s\'"]+', '<PATH>', content)
    content = re.sub(r'/[^\s\'"]+/[^\s\'"]+', '<PATH>', content)
    return content

# wkhtmltopdf management functions
WKHTMLTOPDF_VERSION = '0.12.6'
WKHTMLTOPDF_DOWNLOAD_URL = f'https://github.com/wkhtmltopdf/packaging/releases/download/{WKHTMLTOPDF_VERSION}/wkhtmltox-{WKHTMLTOPDF_VERSION}-1.msvc2015-win64.exe'

def find_wkhtmltopdf() -> str:
    """Find wkhtmltopdf executable, return path or None."""
    # Check common paths
    possible_paths = [
        r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe',
        r'C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe',
        PROJECT_ROOT / 'bin' / 'wkhtmltopdf.exe',  # Portable
    ]
    
    for path in possible_paths:
        path_obj = Path(path) if not isinstance(path, Path) else path
        if path_obj.exists():
            logger.info(f"Found wkhtmltopdf at {path_obj}")
            return str(path_obj)
    
    # Check PATH
    try:
        import shutil
        path_exe = shutil.which('wkhtmltopdf')
        if path_exe:
            logger.info(f"Found wkhtmltopdf in PATH: {path_exe}")
            return path_exe
    except:
        pass
    
    logger.warning("wkhtmltopdf not found")
    return None

def install_wkhtmltopdf_portable() -> str:
    """Download portable version and extract to app folder."""
    try:
        logger.info("Installing portable wkhtmltopdf...")
        portable_url = f'https://github.com/wkhtmltopdf/packaging/releases/download/{WKHTMLTOPDF_VERSION}/wkhtmltox-{WKHTMLTOPDF_VERSION}-1.msvc2015-win64.zip'
        
        import zipfile
        zip_path = tempfile.mktemp(suffix='.zip')
        logger.info(f"Downloading from {portable_url}")
        urllib.request.urlretrieve(portable_url, zip_path)
        
        bin_folder = PROJECT_ROOT / 'bin'
        bin_folder.mkdir(exist_ok=True)
        
        logger.info(f"Extracting to {bin_folder}")
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(bin_folder)
        
        # Find extracted exe
        exe_path = bin_folder / 'bin' / 'wkhtmltopdf.exe'
        if not exe_path.exists():
            # Try alternate structure
            for item in bin_folder.rglob('wkhtmltopdf.exe'):
                exe_path = item
                break
        
        if exe_path.exists():
            logger.info(f"Portable wkhtmltopdf installed at {exe_path}")
            return str(exe_path)
        
        raise Exception("wkhtmltopdf.exe not found after extraction")
    
    except Exception as e:
        logger.error(f"Portable installation failed: {e}", exc_info=True)
        raise

# ============================================================================
# END HELPER FUNCTIONS
# ============================================================================

def render_document_from_file(md_file_path: Path, enable_experimental: bool = False) -> str:
    """Read a document file (markdown or Word), apply feature pipeline, then render HTML."""
    try:
        # Check file size before reading
        file_size = md_file_path.stat().st_size
        if file_size > MAX_FILE_SIZE:
            size_mb = file_size / (1024 * 1024)
            max_mb = MAX_FILE_SIZE / (1024 * 1024)
            logger.warning(f"File too large: {md_file_path} ({size_mb:.2f} MB)")
            return f"<div style='padding: 20px; background: #fee; border: 2px solid #f00; border-radius: 8px;'>"\
                   f"<h2>⚠️ File Too Large</h2>"\
                   f"<p>This file is <strong>{size_mb:.2f} MB</strong>, which exceeds the maximum supported size of <strong>{max_mb:.0f} MB</strong>.</p>"\
                   f"<p><strong>Suggestions:</strong></p>"\
                   f"<ul>"\
                   f"<li>Split the file into smaller documents</li>"\
                   f"<li>Remove unnecessary content or images</li>"\
                   f"<li>Open the file in a text editor instead</li>"\
                   f"</ul></div>"
        
        # Handle Word documents (.docx)
        if md_file_path.suffix.lower() == '.docx':
            if not WORD_INPUT_AVAILABLE:
                logger.error("Word input attempted but mammoth not available")
                return "<p>Word document support not available. Please install mammoth library.</p>"
            
            try:
                html_content = convert_docx_to_html(md_file_path)
                # Process links in the HTML
                html_content = process_links_in_html(html_content, base_path=md_file_path.parent)
                logger.info(f"Successfully rendered Word document: {md_file_path}")
                return html_content
            except Exception as e:
                logger.error(f"Failed to render Word document {md_file_path}: {e}", exc_info=True)
                return f"<p>Error converting Word document: {str(e)}</p>"
        
        # Handle text/markdown files
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
            
        logger.info(f"Rendering document: {md_file_path}, size: {len(md_text)} bytes")
        
    except Exception as e:
        logger.error(f"Error reading file {md_file_path}: {e}", exc_info=True)
        return f"<p>Error reading file: {str(e)}</p>"

    # Apply markdown processing pipeline
    pipeline = FEATURES.build_pipeline(enable_experimental=enable_experimental)
    processed = run_pipeline(md_text, pipeline)
    html_content = render_baseline(processed)
    
    # Process links in the rendered HTML
    html_content = process_links_in_html(html_content, base_path=md_file_path.parent)
    
    return html_content

@app.route('/')
def index():
    """Main page displaying list of all markdown files."""
    md_files = get_markdown_files()
    logger.info(f"Index route: Found {len(md_files)} files in {MD_FOLDER}")
    logger.info(f"PROJECT_ROOT: {PROJECT_ROOT}, Frozen: {getattr(sys, 'frozen', False)}")
    return render_template('index.html', files=md_files, md_folder=str(MD_FOLDER), version=VERSION)

@app.route('/debug/info')
def debug_info():
    """Debug endpoint to show configuration and file discovery."""
    import os
    md_files = get_markdown_files()
    return jsonify({
        'project_root': str(PROJECT_ROOT),
        'md_folder': str(MD_FOLDER),
        'md_folder_exists': MD_FOLDER.exists(),
        'frozen': getattr(sys, 'frozen', False),
        'executable_path': sys.executable,
        'cwd': os.getcwd(),
        'version': VERSION,
        'file_count': len(md_files),
        'files': [{'name': f['name'], 'path': str(f.get('path', 'N/A'))} for f in md_files],
        'config': CONFIG
    })

@app.route('/file/<path:filename>')
def view_file(filename):
    """View a specific markdown file."""
    md_path = Path(MD_FOLDER)
    
    # Smart features temporarily disabled - coming in v1.1/1.2
    # See doc/FUTURE_FEATURES.md for details
    enable_experimental = False  # TODO: Re-enable with proper UI in future release
    
    # Handle both direct filename and relative path
    file_path = None
    
    # Try as relative path first
    potential_path = md_path / filename
    if potential_path.exists() and potential_path.is_file():
        file_path = potential_path
    else:
        # Try adding .md extension
        potential_path = md_path / f"{filename}.md"
        if potential_path.exists():
            file_path = potential_path
        else:
            # Search recursively for matching file
            for f in md_path.rglob('*'):
                if f.is_file() and (f.stem == filename or f.name == filename or str(f.relative_to(md_path)) == filename):
                    file_path = f
                    break
    
    if not file_path or not file_path.exists():
        abort(404)
    
    # Convert to HTML via feature pipeline (baseline + optional experimental)
    html_content = render_document_from_file(file_path, enable_experimental=enable_experimental)
    stat = file_path.stat()
    
    file_info = {
        'name': file_path.stem,
        'filename': file_path.name,
        'content': html_content,
        'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
        'size': f"{stat.st_size / 1024:.2f} KB"
    }
    
    return render_template('view.html', file=file_info)

def get_documentation_files():
    """Get list of available documentation files."""
    if not DOCS_FOLDER.exists():
        return []
    
    docs = []
    for f in DOCS_FOLDER.glob('*.md'):
        if f.name.lower() != 'readme.md':  # Exclude raw readme if user guide exists
             docs.append({
                 'name': f.stem.replace('_', ' ').title(),
                 'filename': f.name,
                 'active': False
             })
    # Sort: User Guide first, then alphabetical
    docs.sort(key=lambda x: (x['filename'] != 'USER_GUIDE.md', x['name']))
    return docs

@app.route('/docs')
@app.route('/docs/<path:filename>')
def documentation(filename=None):
    """Serve the documentation page."""
    if not filename:
        filename = 'USER_GUIDE.md'
    
    # Security: Ensure filename is just a name, not a path
    filename = Path(filename).name
    docs_path = Path(DOCS_FOLDER) / filename
    
    if not docs_path.exists():
        # Fallback to USER_GUIDE if specific file not found, or 404
        if filename != 'USER_GUIDE.md' and (Path(DOCS_FOLDER) / 'USER_GUIDE.md').exists():
             return redirect(url_for('documentation'))
        abort(404, description="Documentation not found")
    
    # Convert documentation markdown to HTML
    html_content = render_document_from_file(docs_path, enable_experimental=False)
    
    # Get navigation list
    nav_items = get_documentation_files()
    for item in nav_items:
        if item['filename'] == filename:
            item['active'] = True
            
    doc_info = {
        'name': filename.replace('.md', '').replace('_', ' ').title(),
        'filename': filename,
        'content': html_content,
        'version': VERSION
    }
    
    return render_template('docs.html', doc=doc_info, nav_items=nav_items)

@app.route('/preview', methods=['POST'])
def preview_file():
    """Preview a document file uploaded from user's filesystem."""
    # Accept file upload via multipart/form-data (better for large files)
    if 'file' in request.files:
        file = request.files['file']
        if not file or file.filename == '':
            abort(400, description="No file provided")
        
        filename = file.filename
        file_ext = Path(filename).suffix.lower()
        
        # Handle Word documents
        if file_ext == '.docx':
            if not WORD_INPUT_AVAILABLE:
                return {
                    "error": "Word Document Support Not Available",
                    "message": "The mammoth library is not installed.",
                    "details": "Please install mammoth to view Word documents."
                }, 503
            
            try:
                # Save to temp file for mammoth processing
                temp_dir = PROJECT_ROOT / 'temp_uploads'
                temp_dir.mkdir(exist_ok=True)
                temp_path = temp_dir / filename
                file.save(temp_path)
                
                # Store in session for link resolution
                session['preview_base_path'] = str(temp_path.parent)
                session['preview_filename'] = filename
                
                # Convert Word to HTML
                html_content = convert_docx_to_html(temp_path)
                html_content = process_links_in_html(html_content, base_path=temp_path.parent, is_preview=True)
                
                file_info = {
                    'name': Path(filename).stem,
                    'filename': filename,
                    'content': html_content,
                    'modified': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'size': f"{temp_path.stat().st_size / 1024:.2f} KB",
                    'preview_mode': True
                }
                
                return render_template('view.html', file=file_info)
                
            except Exception as e:
                logger.error(f"Error processing Word document: {e}", exc_info=True)
                return {
                    "error": "Word Document Conversion Failed",
                    "message": str(e),
                    "filename": filename
                }, 500
        
        # Handle text/markdown files
        try:
            content = file.read().decode('utf-8')
        except UnicodeDecodeError:
            return {
                "error": "Invalid File Encoding",
                "message": "The file must be encoded in UTF-8.",
                "details": "Please save your file as UTF-8 and try again."
            }, 400
    else:
        # Fallback: accept content via form POST (legacy support)
        content = request.form.get('content', '')
        filename = request.form.get('filename', 'Untitled.md')
        if not content:
            abort(400, description="No content provided")
    
    # Check actual content size
    content_size = len(content.encode('utf-8'))
    if content_size > MAX_FILE_SIZE:
        size_mb = content_size / (1024 * 1024)
        max_mb = MAX_FILE_SIZE / (1024 * 1024)
        return {
            "error": "File Too Large",
            "message": f"The file content is {size_mb:.2f} MB, which exceeds the maximum of {max_mb:.0f} MB.",
            "details": "Please split the file into smaller documents or reduce its size.",
            "filename": filename,
            "size_mb": f"{size_mb:.2f}",
            "max_mb": f"{max_mb:.0f}"
        }, 413
    
    # Smart features disabled - using baseline rendering only
    enable_experimental = False
    
    # Use same feature pipeline as file view
    pipeline = FEATURES.build_pipeline(enable_experimental=enable_experimental)
    processed = run_pipeline(content, pipeline)
    html_content = render_baseline(processed)
    html_content = process_links_in_html(html_content, base_path=MD_FOLDER, is_preview=True)
    
    file_info = {
        'name': Path(filename).stem,
        'filename': filename,
        'content': html_content,
        'modified': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'size': f"{len(content) / 1024:.2f} KB",
        'preview_mode': True
    }
    
    return render_template('view.html', file=file_info)

@app.route('/export-pdf', methods=['POST'])
def export_pdf():
    """Export the current document view to PDF using pdfkit/wkhtmltopdf."""
    if not PDF_EXPORT_AVAILABLE:
        return jsonify({
            "error": "PDF export is not available. Please install pdfkit library.",
            "install_guide": "pip install pdfkit"
        }), 503
    
    try:
        data = request.get_json()
        html_content = data.get('html', '')
        filename = data.get('filename', 'document.pdf')
        
        if not html_content:
            abort(400, description="No HTML content provided")
        
        logger.info(f"PDF export requested: {filename}, HTML size: {len(html_content)} bytes")
        
        # Ensure filename has .pdf extension
        if not filename.endswith('.pdf'):
            filename = filename.rsplit('.', 1)[0] + '.pdf' if '.' in filename else filename + '.pdf'
        
        # Try to find wkhtmltopdf executable using new helper
        wkhtmltopdf_path = find_wkhtmltopdf()
        
        if not wkhtmltopdf_path:
            # Not found - offer auto-install
            logger.warning("wkhtmltopdf not found - auto-install available")
            return jsonify({
                'error': 'wkhtmltopdf_not_found',
                'message': 'PDF export requires wkhtmltopdf.',
                'auto_install_available': True,
                'download_size': '15 MB',
                'install_options': ['portable']
            }), 404
        
        # Configure pdfkit with found path
        config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
        
        # PDF options for high quality output
        options = {
            'page-size': 'A4',
            'margin-top': '20mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '20mm',
            'encoding': 'UTF-8',
            'no-outline': None,
            'enable-local-file-access': None,
            'print-media-type': None,
            'dpi': 300,
            'image-quality': 100,
        }
        
        try:
            # Generate PDF from HTML string
            pdf_bytes = pdfkit.from_string(html_content, False, options=options, configuration=config)
            logger.info(f"PDF generated successfully: {len(pdf_bytes)} bytes")
        except OSError as e:
            # wkhtmltopdf execution failed
            if 'No wkhtmltopdf executable found' in str(e) or 'wkhtmltopdf' in str(e).lower():
                logger.error(f"wkhtmltopdf execution failed: {e}")
                return jsonify({
                    'error': 'wkhtmltopdf_not_found',
                    'message': 'wkhtmltopdf could not be executed.',
                    'auto_install_available': True,
                    'details': str(e)
                }), 404
            raise
        
        # Create response with PDF
        response = make_response(pdf_bytes)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        print(f"Error generating PDF: {e}")
        import traceback
        traceback.print_exc()
        abort(500, description=f"Failed to generate PDF: {str(e)}")
        abort(500, description=f"Failed to generate PDF: {str(e)}")

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph in a Word document."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Create bookmark start element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 10000))  # Generate unique ID
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # Create bookmark end element
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    
    # Insert bookmark around the paragraph content
    paragraph._element.insert(0, bookmark_start)
    paragraph._element.append(bookmark_end)

@app.route('/export-word', methods=['POST'])
def export_word():
    """Export the current document view to Word (.docx) format maintaining exact formatting."""
    if not WORD_EXPORT_AVAILABLE:
        return {
            "error": "Word export is not available. Please install htmldocx.",
            "install_guide": "pip install htmldocx"
        }, 503
    
    try:
        # Increase timeout handling for large documents
        data = request.get_json()
        html_content = data.get('html', '')
        filename = data.get('filename', 'document.docx')
        
        if not html_content:
            abort(400, description="No HTML content provided")
        
        # Check HTML content size and provide better feedback
        html_size = len(html_content.encode('utf-8'))
        if html_size > MAX_EXPORT_HTML_SIZE:
            size_mb = html_size / (1024 * 1024)
            max_mb = MAX_EXPORT_HTML_SIZE / (1024 * 1024)
            return {
                "error": "Export Content Too Large",
                "message": f"The HTML content is {size_mb:.2f} MB, exceeding the {max_mb:.0f} MB limit for exports.",
                "details": "Try exporting a smaller section of the document.",
                "tip": "Consider splitting the document into multiple files before exporting."
            }, 413
        
        content_size_mb = html_size / (1024 * 1024)
        print(f"Processing Word export: {filename} ({content_size_mb:.2f} MB)")
        
        if content_size_mb > 30:
            print(f"Warning: Large document detected ({content_size_mb:.2f} MB). Processing may take longer...")
        
        # Ensure filename has .docx extension
        if not filename.endswith('.docx'):
            filename = filename.rsplit('.', 1)[0] + '.docx' if '.' in filename else filename + '.docx'
        
        # Clean HTML content - remove script tags, style tags in head, and other non-document elements
        # Parse HTML and extract only the document content
        # Use lxml parser for better performance on large documents
        from bs4 import BeautifulSoup
        
        try:
            soup = BeautifulSoup(html_content, 'lxml')
        except:
            # Fallback to html.parser if lxml not available
            soup = BeautifulSoup(html_content, 'html.parser')
        
        print("Cleaning HTML content...")
        
        # Remove all script tags
        for script in soup.find_all('script'):
            script.decompose()
        
        # Remove style tags (inline styles in elements will be preserved)
        for style in soup.find_all('style'):
            style.decompose()
        
        # Remove navigation elements
        for nav in soup.find_all('nav'):
            nav.decompose()
        
        # Remove head tag content (but keep the structure for parsing)
        if soup.head:
            # Keep only meta charset for proper encoding
            for tag in soup.head.find_all():
                if tag.name != 'meta' or tag.get('charset') is None:
                    tag.decompose()
        
        print("Extracting main content...")
        
        # Extract only the main content area (markdown-content div)
        main_content = soup.find(class_='markdown-content')
        if main_content:
            # Keep heading IDs for bookmark creation later, but don't add HTML anchors
            # (htmldocx doesn't process them correctly)
            
            # Add inline styles to match rendered version
            # Style tables with MORE EXPLICIT colors for Word
            for table in main_content.find_all('table'):
                # Set table-level styling
                table['style'] = 'border-collapse: collapse; width: 100%; border: 2px solid rgba(99, 102, 241, 0.2); margin-bottom: 20px;'
                table['border'] = '1'
                table['cellpadding'] = '0'
                table['cellspacing'] = '0'
                
                # Find or create thead
                thead = table.find('thead')
                if not thead:
                    # If no thead, check if first row should be header
                    first_row = table.find('tr')
                    if first_row and first_row.find('th'):
                        thead = soup.new_tag('thead')
                        first_row.extract()
                        thead.append(first_row)
                        table.insert(0, thead)
                
                # Style table headers - Word doesn't support gradients, use solid purple
                for th in table.find_all('th'):
                    # Use multiple methods for maximum compatibility
                    th['bgcolor'] = '#6366f1'
                    th['style'] = 'background-color: #6366f1 !important; color: #ffffff !important; padding: 14px 18px; font-weight: 700; text-align: left; border: 1px solid rgba(99, 102, 241, 0.2); font-size: 13px; text-transform: uppercase; letter-spacing: 0.5px;'
                
                # Style table cells with borders and proper padding
                for td in table.find_all('td'):
                    td['style'] = 'padding: 14px 18px; text-align: left; border: 1px solid var(--color-border-default); background-color: #ffffff;'
                
                # Don't apply alternating row colors in HTML - will handle in Word post-processing
                # to ensure proper cell-by-cell styling
            
            # Style headings with STRONG colors to match rendered version
            for h1 in main_content.find_all('h1'):
                h1['style'] = 'color: #6366f1 !important; font-size: 32px; font-weight: 800; margin-top: 24px; margin-bottom: 16px; border-bottom: 3px solid #6366f1; padding-bottom: 8px;'
            
            for h2 in main_content.find_all('h2'):
                h2['style'] = 'color: #8b5cf6 !important; font-size: 24px; font-weight: 700; margin-top: 20px; margin-bottom: 14px; border-bottom: 2px solid #c4b5fd; padding-bottom: 6px;'
            
            for h3 in main_content.find_all('h3'):
                h3['style'] = 'color: #a855f7 !important; font-size: 20px; font-weight: 600; margin-top: 18px; margin-bottom: 12px;'
            
            for h4 in main_content.find_all('h4'):
                h4['style'] = 'color: #c084fc !important; font-size: 18px; font-weight: 600; margin-top: 16px; margin-bottom: 10px;'
            
            for h5 in main_content.find_all('h5'):
                h5['style'] = 'color: #d8b4fe !important; font-size: 16px; font-weight: 600; margin-top: 14px; margin-bottom: 8px;'
            
            for h6 in main_content.find_all('h6'):
                h6['style'] = 'color: #e9d5ff !important; font-size: 14px; font-weight: 600; margin-top: 12px; margin-bottom: 6px;'
            
            # Style code blocks
            for pre in main_content.find_all('pre'):
                pre['style'] = 'background-color: #f5f5f4; padding: 16px; border-radius: 8px; border: 2px solid #e7e5e4; font-family: Consolas, Courier New, monospace; overflow-x: auto; margin: 12px 0;'
            
            # Style inline code
            for code in main_content.find_all('code'):
                if not code.parent or code.parent.name != 'pre':
                    code['style'] = 'background-color: #faf5ff; color: #8b5cf6; padding: 2px 6px; border-radius: 4px; font-family: Consolas, Courier New, monospace; font-size: 0.9em; border: 1px solid #e9d5ff;'
            
            # Style blockquotes
            for blockquote in main_content.find_all('blockquote'):
                blockquote['style'] = 'border-left: 4px solid #6366f1; padding: 16px; margin: 16px 0; background-color: #f5f3ff; color: #44403c;'
            
            # Create a clean HTML structure with just the content
            clean_html = f'<html><head><meta charset="utf-8"></head><body>{str(main_content)}</body></html>'
        else:
            # Fallback: use body content if markdown-content not found
            if soup.body:
                clean_html = f'<html><body>{soup.body.decode_contents()}</body></html>'
            else:
                clean_html = str(soup)
        
        # Create a new Word document
        print("Creating Word document...")
        from docx import Document
        from docx.shared import RGBColor, Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        doc = Document()
        
        # Initialize the HTML to DOCX converter
        print("Converting HTML to Word...")
        new_parser = HtmlToDocx()
        
        # Parse and add cleaned HTML content to the document
        # The htmldocx library maintains most HTML styling including:
        # - Headings (h1-h6) with colors
        # - Paragraphs with formatting
        # - Bold, italic, underline
        # - Tables with colored headers
        # - Lists (ordered and unordered)
        # - Images
        # - Text colors and backgrounds
        # - Font sizes
        # NOTE: htmldocx does NOT support internal anchor links (bookmarks)
        # We need to add these manually after conversion
        new_parser.add_html_to_document(clean_html, doc)
        
        print("Post-processing document...")
        # ==== POST-PROCESSING: Add bookmarks and fix internal hyperlinks ====
        # htmldocx library doesn't support internal document bookmarks/anchors
        # We need to manually add bookmarks to headings and fix TOC links
        
        # Step 1: Find all heading IDs from the original HTML
        heading_ids = {}
        for heading in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            heading_id = heading.get('id')
            heading_text = heading.get_text(strip=True)
            if heading_id and heading_text:
                heading_ids[heading_text] = heading_id
        
        # Step 2: Add bookmarks to headings in the Word document
        for paragraph in doc.paragraphs:
            # Check if this paragraph is a heading with matching text
            if paragraph.style.name.startswith('Heading') or paragraph.text in heading_ids:
                heading_text = paragraph.text.strip()
                if heading_text in heading_ids:
                    bookmark_name = heading_ids[heading_text]
                    # Add bookmark to this paragraph
                    add_bookmark(paragraph, bookmark_name)
        
        # Step 3: Fix internal hyperlinks in the document
        # Find all hyperlinks and convert TOC links from external to internal bookmarks
        for paragraph in doc.paragraphs:
            for hyperlink in paragraph._element.findall('.//w:hyperlink', 
                                                        namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                                                                   'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}):
                # Get the relationship ID
                r_id = hyperlink.get(qn('r:id'))
                if r_id:
                    # Get the target from the relationship
                    try:
                        rel = paragraph.part.rels[r_id]
                        target_url = rel.target_ref
                        
                        # Check if this is an internal anchor link (starts with #)
                        if target_url and target_url.startswith('#'):
                            bookmark_name = target_url[1:]  # Remove the #
                            
                            # Remove the external relationship
                            del paragraph.part.rels[r_id]
                            
                            # Convert to internal bookmark link
                            hyperlink.set(qn('w:anchor'), bookmark_name)
                            # Remove the r:id attribute since it's now an internal link
                            if qn('r:id') in hyperlink.attrib:
                                del hyperlink.attrib[qn('r:id')]
                    except (KeyError, AttributeError):
                        pass  # Skip if relationship not found
        
        # Step 4: Post-process tables in Word document
        # Apply proper formatting that htmldocx might not handle correctly
        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for table in doc.tables:
            # Style table borders
            table.style = 'Table Grid'
            
            # Process header row (first row)
            if len(table.rows) > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    # Set header cell background to purple
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '6366f1')  # Purple background
                    cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Format text: white color, bold, uppercase
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                            run.font.bold = True
                            run.font.size = Pt(11)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Apply alternating row colors for data rows (skip header)
            for idx, row in enumerate(table.rows[1:], start=1):
                for cell in row.cells:
                    # Every other row gets a light background
                    if idx % 2 == 0:
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), 'f9fafb')  # Light gray
                        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Save to BytesIO buffer
        print("Saving document...")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        print(f"Word export complete: {filename}")
        
        # Create response with Word document
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        print(f"Error generating Word document: {e}")
        import traceback
        traceback.print_exc()
        abort(500, description=f"Failed to generate Word document: {str(e)}")

@app.route('/search')
def search():
    """Search through markdown files."""
    query = request.args.get('q', '').strip()
    
    if not query:
        return {'results': []}
    
    query_lower = query.lower()
    results = []
    md_files = get_markdown_files()
    
    for file_info in md_files:
        # Search in filename
        if query_lower in file_info['name'].lower() or query_lower in file_info['filename'].lower():
            results.append({
                'name': file_info['name'],
                'filename': file_info['filename'],
                'path': file_info['relative_path'],
                'folder': file_info['folder'],
                'match_type': 'filename',
                'snippet': f"Found in filename: {file_info['filename']}"
            })
            continue
        
        # Search in file content
        try:
            with open(file_info['path'], 'r', encoding='utf-8') as f:
                content = f.read()
                if query_lower in content.lower():
                    # Find context snippet
                    lines = content.split('\n')
                    for i, line in enumerate(lines):
                        if query_lower in line.lower():
                            # Get context (line before, match line, line after)
                            start = max(0, i - 1)
                            end = min(len(lines), i + 2)
                            snippet = ' '.join(lines[start:end]).strip()
                            if len(snippet) > 150:
                                snippet = snippet[:150] + '...'
                            
                            results.append({
                                'name': file_info['name'],
                                'filename': file_info['filename'],
                                'path': file_info['relative_path'],
                                'folder': file_info['folder'],
                                'match_type': 'content',
                                'snippet': snippet,
                                'line': i + 1
                            })
                            break  # Only show first match per file
        except Exception as e:
            print(f"Error searching file {file_info['path']}: {e}")
            continue
    
    return {'results': results, 'query': query, 'count': len(results)}

@app.route('/static/<path:filename>')
def static_files(filename):
    """Serve static files."""
    return send_from_directory('static', filename)

# ============================================================================
# NEW ROUTES FOR v1.4.0 FEATURES
# ============================================================================

@app.route('/api/get-source/<path:filename>')
def get_source(filename):
    """Get original source content for editing."""
    try:
        file_path = MD_FOLDER / filename
        
        if not file_path.exists() or file_path.suffix.lower() not in ALLOWED_EXTENSIONS:
            logger.warning(f"Source file not found or invalid: {file_path}")
            abort(404)
        
        # Don't allow editing Word files
        if file_path.suffix.lower() == '.docx':
            return jsonify({'error': 'Word documents cannot be edited directly'}), 400
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        logger.info(f"Retrieved source for editing: {filename}")
        return jsonify({'content': content, 'filename': filename, 'size': len(content)})
    
    except Exception as e:
        logger.error(f"Error getting source {filename}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/save-document', methods=['POST'])
def save_document():
    """Save edited document with backup."""
    try:
        data = request.get_json()
        filename = data.get('filename')
        content = data.get('content')
        
        if not filename or content is None:
            return jsonify({'error': 'Missing filename or content'}), 400
        
        # Security: Prevent directory traversal
        if '..' in filename or filename.startswith('/') or filename.startswith('\\'):
            logger.warning(f"Attempted directory traversal: {filename}")
            return jsonify({'error': 'Invalid filename'}), 403
        
        file_path = MD_FOLDER / filename
        
        if not file_path.exists():
            return jsonify({'error': 'File not found'}), 404
        
        # Create backup
        backup_path = file_path.with_suffix(file_path.suffix + '.bak')
        try:
            shutil.copy(file_path, backup_path)
            logger.info(f"Backup created: {backup_path}")
        except Exception as e:
            logger.error(f"Failed to create backup: {e}")
            return jsonify({'error': 'Failed to create backup'}), 500
        
        # Write new content
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"Document saved: {filename}, size: {len(content)} bytes")
            return jsonify({'success': True, 'backup': str(backup_path), 'size': len(content)})
        except Exception as e:
            # Restore from backup if write failed
            if backup_path.exists():
                shutil.copy(backup_path, file_path)
            logger.error(f"Failed to save document: {e}")
            return jsonify({'error': 'Failed to save document'}), 500
    
    except Exception as e:
        logger.error(f"Error in save_document: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-logs')
def download_logs():
    """Create sanitized ZIP of logs for user download."""
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f'omnidoc_logs_{timestamp}.zip'
        zip_path = LOG_DIR / zip_filename
        
        logger.info(f"Creating log archive: {zip_filename}")
        
        # Create ZIP
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for log_file in LOG_DIR.glob('omnidoc.log*'):
                try:
                    # Sanitize content
                    sanitized_content = sanitize_log_content(log_file.read_text(encoding='utf-8', errors='ignore'))
                    zipf.writestr(log_file.name, sanitized_content)
                except Exception as e:
                    logger.error(f"Error adding log file {log_file}: {e}")
        
        logger.info(f"Log archive created: {zip_path}")
        return send_from_directory(LOG_DIR, zip_filename, as_attachment=True)
    
    except Exception as e:
        logger.error(f"Error creating log archive: {e}", exc_info=True)
        return jsonify({'error': 'Failed to create log archive'}), 500

@app.route('/api/workspaces', methods=['GET'])
def get_workspaces():
    """Get list of configured workspaces."""
    try:
        config = load_config()
        return jsonify({
            'workspaces': config['workspaces'],
            'active': config['active_workspace'],
            'recent': config.get('recent_workspaces', [])
        })
    except Exception as e:
        logger.error(f"Error getting workspaces: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/workspaces', methods=['POST'])
def add_workspace():
    """Add new workspace directory."""
    try:
        data = request.get_json()
        workspace_path = data.get('path')
        
        if not workspace_path:
            return jsonify({'error': 'Missing workspace path'}), 400
        
        path_obj = Path(workspace_path)
        
        if not path_obj.exists():
            return jsonify({'error': 'Directory does not exist'}), 400
        
        if not path_obj.is_dir():
            return jsonify({'error': 'Path is not a directory'}), 400
        
        # Security checks
        if not is_safe_workspace(path_obj):
            logger.warning(f"Unsafe workspace rejected: {workspace_path}")
            return jsonify({'error': 'Access to this directory is not allowed for security reasons'}), 403
        
        # Add to config
        config = load_config()
        workspace_str = str(path_obj.resolve())
        
        if workspace_str not in config['workspaces']:
            config['workspaces'].append(workspace_str)
            save_config(config)
            logger.info(f"Workspace added: {workspace_str}")
        
        return jsonify({'success': True, 'workspace': workspace_str})
    
    except Exception as e:
        logger.error(f"Error adding workspace: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/workspaces/active', methods=['POST'])
def set_active_workspace():
    """Switch active workspace."""
    try:
        data = request.get_json()
        workspace_path = data.get('path')
        
        if not workspace_path:
            return jsonify({'error': 'Missing workspace path'}), 400
        
        # Resolve path to ensure it matches stored format
        path_obj = Path(workspace_path).resolve()
        workspace_str = str(path_obj)
        
        config = load_config()
        
        if workspace_str not in config['workspaces']:
            return jsonify({'error': f'Workspace not configured: {workspace_str}'}), 400
        
        config['active_workspace'] = workspace_str
        
        # Update recent list
        if 'recent_workspaces' not in config:
            config['recent_workspaces'] = []
        
        if workspace_path in config['recent_workspaces']:
            config['recent_workspaces'].remove(workspace_path)
        
        config['recent_workspaces'].insert(0, workspace_path)
        config['recent_workspaces'] = config['recent_workspaces'][:5]  # Keep last 5
        
        save_config(config)
        
        # Update global MD_FOLDER
        global MD_FOLDER
        MD_FOLDER = Path(workspace_path)
        
        logger.info(f"Active workspace changed to: {workspace_path}")
        return jsonify({'success': True, 'active_workspace': workspace_path})
    
    except Exception as e:
        logger.error(f"Error setting active workspace: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/workspaces/<path:workspace_path>', methods=['DELETE'])
def delete_workspace(workspace_path):
    """Remove workspace from configuration."""
    try:
        config = load_config()
        
        if workspace_path in config['workspaces']:
            # Prevent deleting active workspace
            if workspace_path == config['active_workspace']:
                return jsonify({'error': 'Cannot delete active workspace'}), 400
            
            config['workspaces'].remove(workspace_path)
            save_config(config)
            logger.info(f"Workspace removed: {workspace_path}")
            return jsonify({'success': True})
        
        return jsonify({'error': 'Workspace not found'}), 404
    
    except Exception as e:
        logger.error(f"Error deleting workspace: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/install-wkhtmltopdf', methods=['POST'])
def api_install_wkhtmltopdf():
    """Endpoint to trigger wkhtmltopdf installation."""
    try:
        data = request.get_json()
        mode = data.get('mode', 'portable')  # 'portable' or 'system'
        
        logger.info(f"wkhtmltopdf installation requested: mode={mode}")
        
        if mode == 'portable':
            exe_path = install_wkhtmltopdf_portable()
            if exe_path:
                return jsonify({'success': True, 'path': str(exe_path)})
        else:
            # System install requires admin privileges - not implemented
            return jsonify({'success': False, 'error': 'System installation not implemented'}), 501
        
        return jsonify({'success': False, 'error': 'Installation failed'}), 500
    
    except Exception as e:
        logger.error(f"Installation error: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/browse-folder', methods=['GET'])
def browse_folder():
    """Open native folder browser dialog."""
    try:
        import tkinter as tk
        from tkinter import filedialog
        
        # Create and hide root window
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        
        # Open folder dialog
        folder_path = filedialog.askdirectory(
            title='Select Workspace Folder',
            mustexist=True
        )
        
        root.destroy()
        
        if folder_path:
            return jsonify({'success': True, 'path': folder_path})
        else:
            return jsonify({'success': False, 'message': 'No folder selected'}), 400
            
    except ImportError:
        logger.warning("tkinter not available for native folder browser")
        return jsonify({'error': 'Native folder browser not available. Please enter path manually.'}), 501
    except Exception as e:
        logger.error(f"Error opening folder browser: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

# ============================================================================
# END NEW ROUTES
# ============================================================================

if __name__ == '__main__':
    app.run(debug=True, host='localhost', port=8000)
